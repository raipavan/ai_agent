from __future__ import annotations

import io
import json
import os
import re
import time
import uuid
import wave
from pathlib import Path

from fastapi import APIRouter, Request, HTTPException, UploadFile, File, Depends
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from loguru import logger

from core.auth import get_current_user, get_current_user_optional, _decode_jwt
from core.state import get_state, save_role_state, normalize_console_role, resolved_greeting_text, _CAMPAIGN_DATA
from core.utils import range_file_response
from config import settings
from core.outbound_numbers import resolve_outbound_from_number
from core.vobiz_credentials import resolve_vobiz_credentials


def _role_from_jwt(request: Request) -> str | None:
    """Extract role from JWT Authorization header, or None."""
    from core.auth import dashboard_role_for_token

    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        payload = _decode_jwt(auth[7:])
        if payload and payload.get("role"):
            return dashboard_role_for_token(
                payload.get("email"),
                payload.get("role"),
            )
    return None


def _role_from_request(request: Request, default: str = "sales_1") -> str:
    """JWT role, or ``?role=sales_1|sales_2`` when the console role toggle is used."""
    from core.auth import console_role_from_request
    from core.state import normalize_console_role

    role_param = (request.query_params.get("role") or "").strip()
    if role_param:
        return normalize_console_role(role_param)
    return console_role_from_request(request, default=default)
from core.phone_norm import norm_phone_str
from core.greeting_pcm import load_recorded_greeting_pcm
from core.storage import insert_manual_call, mark_manual_call_failed
from core.utils import _build_opening_line
from core.worker import _prime_opening_audio
from services.call_recording import resolve_session_recording_path
from services.vobiz_bridge import make_vobiz_call

router = APIRouter(tags=["console"])


def _readable_transcript_lines(raw: str) -> tuple[str, list[str]]:
    """Return (joined readable text, list of lines) from JSONL or plain text."""
    lines_out: list[str] = []
    for line in (raw or "").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
            role = obj.get("role") or obj.get("type", "")
            content = obj.get("content") or obj.get("text") or obj.get("message", "")
            if role in ("user", "assistant") and content:
                lines_out.append(f"{role.capitalize()}: {content.strip()}")
        except Exception:
            lines_out.append(line)
    if lines_out:
        return "\n".join(lines_out), lines_out
    return (raw or "").strip(), []


def _recommended_actions_from_analysis(analysis: dict) -> list[str]:
    bullets: list[str] = []
    disp = (analysis.get("disposition") or "").strip()
    if disp:
        bullets.append(f"Disposition: {disp}")
    ns = analysis.get("next_steps")
    if isinstance(ns, list):
        for x in ns:
            s = str(x).strip().lstrip("•-*").strip()
            if s:
                bullets.append(s)
        return bullets[:24]
    text = str(ns or "").strip()
    if text:
        parts = [p.strip().lstrip("•-*").strip() for p in re.split(r"[\n;]+", text)]
        parts = [p for p in parts if p]
        if parts:
            bullets.extend(parts)
        else:
            bullets.append(text)
    return bullets[:24]


def _manual_call_row_to_summary(row: dict) -> dict:
    log_id = row.get("log_id") or row.get("camp_id") or ""
    recording_available = bool(log_id and resolve_session_recording_path(log_id))
    return {
        "id": row["id"],
        "role": row["role"],
        "camp_id": row["camp_id"],
        "to_phone": row["to_phone"],
        "callee_name": row["callee_name"],
        "status": row["status"],
        "started_at": row.get("started_at"),
        "ended_at": row.get("ended_at"),
        "duration_sec": row.get("duration_sec"),
        "disposition": row.get("disposition") or "",
        "summary": (row.get("summary") or "")[:400],
        "recording_available": recording_available,
        "recording_url": f"/api/manual/calls/{row['id']}/recording?role={row['role']}" if recording_available else "",
    }


def _manual_call_detail_response(row: dict) -> dict:
    from core.worker import _read_transcript_jsonl

    role = row["role"]
    log_id = row.get("log_id") or row.get("camp_id") or ""
    raw = _read_transcript_jsonl(role, log_id) if log_id else ""
    readable, line_list = _readable_transcript_lines(raw)
    aj: dict = {}
    try:
        if (row.get("analysis_json") or "").strip():
            parsed = json.loads(row["analysis_json"])
            if isinstance(parsed, dict):
                aj = parsed
    except Exception:
        aj = {}
    # Fall back to the post-call analysis transcript when no JSONL live
    # transcript exists (WS-bridged calls store it only in analysis_json).
    if not readable and (aj.get("transcript") or "").strip():
        readable = str(aj["transcript"]).strip()
        line_list = [{"speaker": "All", "text": readable}]
    recording_available = False
    if (log_id or "").strip():
        recording_available = bool(resolve_session_recording_path(log_id))
    # Prefer flattened columns when present
    if not aj.get("summary") and row.get("summary"):
        aj = {**aj, "summary": row.get("summary")}
    if not aj.get("disposition") and row.get("disposition"):
        aj = {**aj, "disposition": row.get("disposition")}
    if not aj.get("next_steps") and row.get("next_steps"):
        aj = {**aj, "next_steps": row.get("next_steps")}
    if "next_action" not in aj:
        aj = {**aj, "next_action": None}
    if not aj.get("emotion_label") and row.get("emotion_label"):
        aj = {
            **aj,
            "emotion_label": row.get("emotion_label"),
            "emotion_rationale": row.get("emotion_rationale"),
            "emotion_confidence": row.get("emotion_confidence"),
        }
    return {
        "id": row["id"],
        "role": row["role"],
        "camp_id": row["camp_id"],
        "to_phone": row["to_phone"],
        "callee_name": row["callee_name"],
        "status": row["status"],
        "started_at": row.get("started_at"),
        "ended_at": row.get("ended_at"),
        "duration_sec": row.get("duration_sec"),
        "log_id": log_id,
        "transcript_raw": raw,
        "transcript_readable": readable,
        "transcript_lines": line_list,
        "summary": aj.get("summary")
        or row.get("summary")
        or ((row.get("error") or "") if (row.get("status") or "") == "failed" else "")
        or "",
        "disposition": aj.get("disposition") or row.get("disposition") or "",
        "next_steps": aj.get("next_steps") or row.get("next_steps") or "",
        "emotion_label": aj.get("emotion_label") or row.get("emotion_label") or "",
        "emotion_rationale": aj.get("emotion_rationale") or row.get("emotion_rationale") or "",
        "emotion_confidence": aj.get("emotion_confidence", row.get("emotion_confidence")),
        "recommended_actions": _recommended_actions_from_analysis(aj),
        "rating": aj.get("rating"),
        "analysis": aj,
        "error": row.get("error") or "",
        "recording_available": recording_available,
        "recording_url": f"/api/manual/calls/{row['id']}/recording?role={row['role']}" if recording_available else "",
    }


def _incoming_call_row_to_summary(row: dict) -> dict:
    log_id = row.get("log_id") or row.get("camp_id") or ""
    recording_available = bool(log_id and resolve_session_recording_path(log_id))
    return {
        "id": row["id"],
        "role": row["role"],
        "camp_id": row["camp_id"],
        "from_phone": row["from_phone"],
        "caller_name": row.get("caller_name") or "",
        "callee_name": row.get("caller_name") or "",
        "status": row["status"],
        "started_at": row.get("started_at"),
        "ended_at": row.get("ended_at"),
        "duration_sec": row.get("duration_sec"),
        "disposition": row.get("disposition") or "",
        "summary": (row.get("summary") or "")[:400],
        "recording_available": recording_available,
        "recording_url": f"/api/incoming/calls/{row['id']}/recording?role={row['role']}" if recording_available else "",
    }


def _incoming_call_detail_response(row: dict) -> dict:
    from core.worker import _read_transcript_jsonl

    role = row["role"]
    log_id = row.get("log_id") or row.get("camp_id") or ""
    raw = _read_transcript_jsonl(role, log_id) if log_id else ""
    readable, line_list = _readable_transcript_lines(raw)
    aj: dict = {}
    try:
        if (row.get("analysis_json") or "").strip():
            parsed = json.loads(row["analysis_json"])
            if isinstance(parsed, dict):
                aj = parsed
    except Exception:
        aj = {}
    # Fall back to the post-call analysis transcript when no JSONL live
    # transcript exists (WS-bridged calls store it only in analysis_json).
    if not readable and (aj.get("transcript") or "").strip():
        readable = str(aj["transcript"]).strip()
        line_list = [{"speaker": "All", "text": readable}]
    recording_available = False
    if (log_id or "").strip():
        recording_available = bool(resolve_session_recording_path(log_id))
    if not aj.get("summary") and row.get("summary"):
        aj = {**aj, "summary": row.get("summary")}
    if not aj.get("disposition") and row.get("disposition"):
        aj = {**aj, "disposition": row.get("disposition")}
    if not aj.get("next_steps") and row.get("next_steps"):
        aj = {**aj, "next_steps": row.get("next_steps")}
    if "next_action" not in aj:
        aj = {**aj, "next_action": None}
    if not aj.get("emotion_label") and row.get("emotion_label"):
        aj = {
            **aj,
            "emotion_label": row.get("emotion_label"),
            "emotion_rationale": row.get("emotion_rationale"),
            "emotion_confidence": row.get("emotion_confidence"),
        }
    return {
        "id": row["id"],
        "role": row["role"],
        "camp_id": row["camp_id"],
        "from_phone": row["from_phone"],
        "caller_name": row["caller_name"],
        "callee_name": row.get("caller_name") or "",
        "lead_name": row.get("caller_name") or "",
        "status": row["status"],
        "started_at": row.get("started_at"),
        "ended_at": row.get("ended_at"),
        "duration_sec": row.get("duration_sec"),
        "log_id": log_id,
        "transcript_raw": raw,
        "transcript_readable": readable,
        "transcript_lines": line_list,
        "summary": aj.get("summary")
        or row.get("summary")
        or ((row.get("error") or "") if (row.get("status") or "") == "failed" else "")
        or "",
        "disposition": aj.get("disposition") or row.get("disposition") or "",
        "next_steps": aj.get("next_steps") or row.get("next_steps") or "",
        "emotion_label": aj.get("emotion_label") or row.get("emotion_label") or "",
        "emotion_rationale": aj.get("emotion_rationale") or row.get("emotion_rationale") or "",
        "emotion_confidence": aj.get("emotion_confidence", row.get("emotion_confidence")),
        "recommended_actions": _recommended_actions_from_analysis(aj),
        "rating": aj.get("rating"),
        "analysis": aj,
        "error": row.get("error") or "",
        "recording_available": recording_available,
        "recording_url": f"/api/incoming/calls/{row['id']}/recording?role={row['role']}" if recording_available else "",
    }


def _pcm_s16le_to_wav(pcm: bytes, sample_rate: int) -> bytes:
    buf = io.BytesIO()
    with wave.open(buf, "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        wf.writeframes(pcm)
    return buf.getvalue()


@router.get("/api/tuning")
async def get_tuning(request: Request):
    role = _role_from_request(request)
    logger.info("get_tuning role={!r} qparams={}", role, dict(request.query_params))
    state = get_state(role)

    from core.opening_line import packaged_fallback_greeting
    from core.role_sandbox import coerce_role_prompt, coerce_role_rag, coerce_stored_greeting
    from prompts.priya import get_role_prompt_text, get_role_rag_source_text

    file_prompt = get_role_prompt_text(role)
    file_rag = get_role_rag_source_text(role)
    prompt = coerce_role_prompt(role, state.get("prompt", ""), file_prompt)
    rag = coerce_role_rag(role, state.get("rag", ""), file_rag)
    gt = coerce_stored_greeting(role, state.get("greeting_text"))
    greeting = gt if gt else packaged_fallback_greeting(role)

    # Prompt Management page — named sections stored in role_state.prompt_parts.
    try:
        parts = dict(state.get("prompt_parts") or {})
    except Exception:
        parts = {}
    prompts = {
        "system": parts.get("system") or prompt or file_prompt,
        "greeting": parts.get("greeting") or greeting,
        "fallback": parts.get("fallback") or "",
        "escalation": parts.get("escalation") or "",
        "closing": parts.get("closing") or "",
        "compliance": parts.get("compliance") or "",
        "variables": parts.get("variables") or "",
        "version": parts.get("version") or "v3.2",
        "published": bool(parts.get("published")),
        "history": parts.get("history") or [],
    }

    return {
        "role": role,
        "prompt": prompt,
        "rag": rag,
        "greeting_text": greeting,
        "prompts": prompts,
    }

class TuningUpdate(BaseModel):
    prompt: str = ""
    rag: str = ""
    greeting_text: str = ""

@router.post("/api/tuning")
async def update_tuning(data: TuningUpdate, request: Request):
    role = _role_from_request(request)
    logger.info("update_tuning role={!r} qparams={}", role, dict(request.query_params))

    if role not in ("sales_1", "sales_2"):
        logger.warning("update_tuning received invalid role={!r}", role)
        raise HTTPException(400, f"Invalid role: {role}")

    from core.greeting_text_utils import coerce_stored_greeting
    from core.role_sandbox import validate_role_tuning

    tuning_err = validate_role_tuning(
        role,
        prompt=data.prompt or "",
        rag=data.rag or "",
        greeting=data.greeting_text or "",
    )
    if tuning_err:
        raise HTTPException(400, tuning_err)

    greeting_out = coerce_stored_greeting(role, data.greeting_text or "")
    save_role_state(role, prompt=data.prompt, rag=data.rag, greeting_text=greeting_out)

    # Keep prompt + KB files in sync — build_role_system_prompt() prefers non-empty DB,
    # then falls back to these files when the DB field is empty.
    from prompts.priya import set_role_prompt_text, set_role_rag_source_text

    set_role_prompt_text(role, data.prompt)
    set_role_rag_source_text(role, data.rag)

    # Keep the chunked RAG index fresh after a KB save. Indexing failures must
    # never fail the tuning save itself.
    try:
        from rag import index_role_rag

        index_role_rag(role, data.rag)
    except Exception as exc:
        logger.warning("RAG chunk index refresh failed for role={}: {}", role, exc)

    try:
        from core.notifications import push_notification

        push_notification(
            role,
            "Prompt & knowledge updated",
            f"System prompt and RAG saved for {role}",
            kind="system",
        )
    except Exception as ne:
        logger.warning("Failed to push tuning notification: {}", ne)

    return {"status": "ok", "saved_role": role}


class PromptFieldUpdate(BaseModel):
    field: str = ""
    value: str = ""


_ALLOWED_PROMPT_FIELDS = frozenset(
    {"system", "greeting", "fallback", "escalation", "closing", "compliance", "variables"}
)


@router.post("/api/tuning/prompt")
async def update_prompt_field(data: PromptFieldUpdate, request: Request):
    """Save one named prompt section (Prompt Management page) for the role."""
    role = _role_from_request(request)
    if role not in ("sales_1", "sales_2"):
        raise HTTPException(400, f"Invalid role: {role}")

    field = (data.field or "").strip().lower()
    if field not in _ALLOWED_PROMPT_FIELDS:
        raise HTTPException(400, f"Unknown prompt field: {field}")

    state = get_state(role)
    try:
        parts = dict(state.get("prompt_parts") or {})
    except Exception:
        parts = {}
    parts[field] = data.value or ""

    # Bump version + prepend a history entry (keep last 20).
    import re as _re

    ver = str(parts.get("version") or "v3.2")
    m = _re.search(r"(\d+)$", ver)
    n = (int(m.group(1)) + 1) if m else 2
    parts["version"] = f"v3.{n}"
    from datetime import datetime, timezone, timedelta

    now = datetime.now(timezone(timedelta(hours=5, minutes=30)))
    hist = parts.get("history") or []
    hist.insert(
        0,
        {
            "version": parts["version"],
            "date": now.strftime("%Y-%m-%d %H:%M"),
            "changes": f"{field} updated",
            "status": "Published" if parts.get("published") else "Draft",
        },
    )
    parts["history"] = hist[:20]

    # System prompt and greeting are the two fields that feed the live call.
    prompt_val = data.value if field == "system" else None
    greeting_val = None
    if field == "greeting":
        from core.greeting_text_utils import coerce_stored_greeting

        greeting_val = coerce_stored_greeting(role, data.value)
    save_role_state(role, prompt=prompt_val, greeting_text=greeting_val, prompt_parts=parts)

    if field == "system":
        from prompts.priya import set_role_prompt_text

        set_role_prompt_text(role, data.value)

    try:
        from core.notifications import push_notification

        push_notification(
            role,
            "Prompt updated",
            f"{field} prompt saved for {role} (v{parts['version'][2:]})",
            kind="system",
        )
    except Exception as ne:
        logger.warning("Failed to push prompt-update notification: {}", ne)

    return {"status": "ok", "field": field, "role": role, "version": parts["version"]}


@router.post("/api/tuning/publish")
async def publish_prompt(request: Request):
    """Mark the role's prompt set as published (Prompt Management page)."""
    role = _role_from_request(request)
    if role not in ("sales_1", "sales_2"):
        raise HTTPException(400, f"Invalid role: {role}")

    state = get_state(role)
    try:
        parts = dict(state.get("prompt_parts") or {})
    except Exception:
        parts = {}
    parts["published"] = True
    from datetime import datetime, timezone, timedelta

    now = datetime.now(timezone(timedelta(hours=5, minutes=30)))
    hist = parts.get("history") or []
    hist.insert(
        0,
        {
            "version": parts.get("version") or "v3.2",
            "date": now.strftime("%Y-%m-%d %H:%M"),
            "changes": "Published",
            "status": "Published",
        },
    )
    parts["history"] = hist[:20]
    save_role_state(role, prompt_parts=parts)

    try:
        from core.notifications import push_notification

        push_notification(role, "Prompt published", f"Prompt set published for {role}", kind="system")
    except Exception as ne:
        logger.warning("Failed to push prompt-publish notification: {}", ne)

    return {"status": "ok", "role": role, "published": True}


class GreetingTextBody(BaseModel):
    greeting_text: str = ""


@router.post("/api/tuning/record-greeting")
async def record_greeting_tts(data: GreetingTextBody, request: Request):
    """Capture PCM for greeting_{role}.pcm using Gemini Live (native call voice)."""
    role = _role_from_request(request)
    text = (data.greeting_text or "").strip()
    if not text:
        raise HTTPException(400, "greeting_text is required")

    from config import settings
    from core.greeting_pcm import _generate_and_cache_greeting, greeting_pcm_paths

    try:
        result = await _generate_and_cache_greeting(
            role,
            text,
            settings.gemini_live_voice,
        )
    except Exception as exc:
        logger.exception("record-greeting failed")
        raise HTTPException(503, f"Greeting generation failed: {exc}") from exc

    if not result:
        raise HTTPException(503, "Greeting generation failed")

    pcm, sr = result
    out_path, _ = greeting_pcm_paths(role)

    return {
        "status": "ok",
        "path": str(out_path),
        "bytes": len(pcm),
        "sample_rate": sr,
        "engine": "live",
    }


@router.post("/api/tuning/capture-greeting-live")
async def capture_greeting_live(data: GreetingTextBody, request: Request):
    """Capture opening audio from Gemini Live (native voice) and save greeting_{role}.pcm.

    Returns WAV for immediate playback; PCM on disk is what calls use before Live connects.
    Query ``variant=inbound`` saves ``greeting_{role}_inbound.pcm`` (inbound DID legs).
    """
    role = _role_from_request(request)
    variant = (request.query_params.get("variant") or "").strip().lower()
    text = (data.greeting_text or "").strip()
    if not text:
        raise HTTPException(400, "greeting_text is required")

    from services.live_greeting_capture import capture_live_greeting_pcm, save_greeting_pcm_file

    logger.info(
        "capture-greeting-live: role={} variant={} text_len={}",
        role,
        variant or "(default)",
        len(text),
    )

    try:
        pcm, sr = await capture_live_greeting_pcm(role, text)
        if variant:
            out_path = save_greeting_pcm_file(
                role, pcm, sr, variant=variant, greeting_text=text
            )
        else:
            from config import settings
            from core.greeting_pcm import _write_greeting_cache_files, greeting_pcm_paths

            live_voice = (settings.gemini_live_voice or "Leda").strip()
            _write_greeting_cache_files(
                role, text, pcm, sr, source="gemini_live_capture", voice=live_voice
            )
            out_path, _ = greeting_pcm_paths(role)
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    except RuntimeError as exc:
        logger.warning("capture-greeting-live failed role={}: {}", role, exc)
        raise HTTPException(503, str(exc)) from exc
    except Exception as exc:
        logger.exception("capture-greeting-live failed role={}", role)
        raise HTTPException(503, f"Live capture failed: {exc}") from exc

    wav = _pcm_s16le_to_wav(pcm, sr)
    return Response(
        content=wav,
        media_type="audio/wav",
        headers={
            "Cache-Control": "no-store",
            "X-Sample-Rate": str(sr),
            "X-Role": role,
            "X-Greeting-Bytes": str(len(pcm)),
            "X-Greeting-Path": str(out_path),
            "X-Greeting-Source": "gemini_live",
        },
    )


@router.post("/api/tuning/upload-doc")
async def upload_doc(request: Request, file: UploadFile = File(...)):
    role = _role_from_request(request)
    # extract text
    content = await file.read()
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith(".txt"):
            text = content.decode("utf-8", errors="replace")
        elif filename.endswith(".pdf"):
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            import docx
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            raise HTTPException(400, "Unsupported file type")
    except Exception as e:
        logger.error(f"Failed to extract document: {e}")
        raise HTTPException(500, f"Extraction failed: {e}")
        
    from prompts.priya import get_role_rag_source_text, set_role_rag_source_text

    state = get_state(role)
    current_rag = get_role_rag_source_text(role) or state.get("rag", "")
    new_rag = current_rag + "\n\n" + text if current_rag else text
    save_role_state(role, rag=new_rag)
    set_role_rag_source_text(role, new_rag)

    # Rebuild the chunk index for the appended document. Failures only log —
    # the upload itself already succeeded.
    try:
        from rag import index_role_rag

        index_role_rag(role, new_rag)
    except Exception as exc:
        logger.warning("RAG chunk index refresh failed after upload for role={}: {}", role, exc)

    return {"status": "ok", "filename": file.filename, "extracted_length": len(text)}


# ── RAG / Knowledge Base management (UI-driven) ─────────────────────────


def _current_role_rag(role: str) -> str:
    """Effective RAG text for a role: DB value first, packaged files as default."""
    from core.role_sandbox import coerce_role_rag
    from prompts.priya import get_role_rag_source_text

    state = get_state(role)
    return coerce_role_rag(role, state.get("rag", ""), get_role_rag_source_text(role))


@router.get("/api/rag")
async def get_rag(request: Request):
    """Full RAG knowledge text + chunk stats for one role (console editor)."""
    role = _role_from_request(request)
    rag = _current_role_rag(role)
    from rag import role_chunk_count

    return {
        "role": role,
        "rag": rag,
        "chunks": role_chunk_count(role),
        "sources": 1,
        "indexed": role_chunk_count(role) > 0,
        "engine": "keyword-fts-chunks",
    }


class RagUpdate(BaseModel):
    rag: str = ""


@router.post("/api/rag")
async def update_rag(data: RagUpdate, request: Request):
    """Save the role's RAG knowledge from the console and rebuild its chunk index."""
    role = _role_from_request(request)
    if role not in ("sales_1", "sales_2"):
        raise HTTPException(400, f"Invalid role: {role}")

    save_role_state(role, rag=data.rag or "")
    from prompts.priya import set_role_rag_source_text

    set_role_rag_source_text(role, data.rag or "")

    chunks = 0
    try:
        from rag import index_role_rag

        chunks = index_role_rag(role, data.rag or "")
    except Exception as exc:
        logger.warning("RAG chunk index refresh failed for role={}: {}", role, exc)

    return {"status": "ok", "saved_role": role, "chunks": chunks}


@router.post("/api/rag/reindex")
async def reindex_rag(request: Request):
    """Rebuild the chunk index for a role from its currently saved RAG text."""
    role = _role_from_request(request)
    rag = _current_role_rag(role)
    from rag import index_role_rag

    chunks = index_role_rag(role, rag)
    return {"status": "ok", "role": role, "chunks": chunks}


@router.get("/api/rag/query")
async def query_rag(request: Request, q: str = ""):
    """Live chunk retrieval preview for the console search box."""
    role = _role_from_request(request)
    from rag import query_role_chunks

    results = query_role_chunks(role, (q or "").strip())
    return {"role": role, "query": q, "results": results}


@router.post("/api/rag/upload")
async def upload_rag_file(request: Request, file: UploadFile = File(...)):
    """Append a document's text to the role's RAG knowledge (UI upload)."""
    role = _role_from_request(request)
    content = await file.read()
    filename = (file.filename or "").lower()
    text = ""
    try:
        if filename.endswith(".txt") or filename.endswith(".md"):
            text = content.decode("utf-8", errors="replace")
        elif filename.endswith(".csv"):
            text = content.decode("utf-8", errors="replace")
        elif filename.endswith(".pdf"):
            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        elif filename.endswith(".docx"):
            import docx

            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            raise HTTPException(400, "Unsupported file type (use .txt, .md, .csv, .pdf, .docx)")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Failed to extract document {}: {}", file.filename, exc)
        raise HTTPException(500, f"Extraction failed: {exc}")

    current_rag = _current_role_rag(role)
    new_rag = (current_rag + "\n\n" + text).strip() if current_rag.strip() else text.strip()
    save_role_state(role, rag=new_rag)
    from prompts.priya import set_role_rag_source_text

    set_role_rag_source_text(role, new_rag)

    chunks = 0
    try:
        from rag import index_role_rag

        chunks = index_role_rag(role, new_rag)
    except Exception as exc:
        logger.warning("RAG chunk index refresh failed after upload for role={}: {}", role, exc)

    return {
        "status": "ok",
        "filename": file.filename,
        "extracted_length": len(text),
        "chunks": chunks,
    }

class VobizUpdate(BaseModel):
    auth_id: str
    auth_token: str
    from_number: str
    public_url: str

@router.post("/api/settings/vobiz")
async def update_vobiz(data: VobizUpdate, request: Request):
    role = _role_from_request(request)
    state = get_state(role)
    vobiz_config = state.get("vobiz", {})
    vobiz_config.update({
        "auth_id": data.auth_id,
        "auth_token": data.auth_token,
        "from_number": data.from_number,
        "public_url": data.public_url
    })
    save_role_state(role, vobiz_config=vobiz_config)
    return {"status": "ok"}

class ManualCallReq(BaseModel):
    to: str
    callee_name: str

@router.post("/api/manual/call")
async def manual_call(
    data: ManualCallReq,
    request: Request,
    _user: dict = Depends(get_current_user_optional),
):
    role = _role_from_request(request)
    state = get_state(role)
    vobiz_config = state.get("vobiz", {})

    auth_id, auth_token, from_number, v_base = resolve_vobiz_credentials(role, vobiz_config)

    # Use round-robin phone number selection for sales roles
    if role in ("sales_1", "sales_2"):
        from core.worker import get_next_phone_number
        from_number = get_next_phone_number(role, vobiz_config)
        logger.info(f"Manual call: Using phone number {from_number} for {role} (round-robin)")

    if not auth_id or not auth_token:
        raise HTTPException(400, "Vobiz credentials not configured")
    if not from_number.strip():
        raise HTTPException(400, "Outbound caller ID (from_number) is not configured for this role")
    if not v_base:
        raise HTTPException(400, "VOBIZ_PUBLIC_BASE_URL is not configured")

    to_norm = norm_phone_str((data.to or "").strip())
    if not to_norm:
        raise HTTPException(400, "Invalid phone number — enter 10 digits (after +91), or a full number starting with + (e.g. +971…).")

    camp_id = f"manual_{role}_{uuid.uuid4()}"
    manual_row: dict = {
        "_role": role,
        "_manual_leg": True,
        "phone": to_norm,
        "name": (data.callee_name or "").strip() or "Unknown",
    }
    from core.greeting_text_utils import coerce_stored_greeting

    gt = coerce_stored_greeting(role, (state.get("greeting_text") or "").strip())
    if gt:
        manual_row["greeting_text"] = gt

    manual_call_id = await insert_manual_call(
        role,
        camp_id,
        to_norm,
        (data.callee_name or "").strip() or "Unknown",
    )
    _CAMPAIGN_DATA[camp_id] = manual_row

    opening_text = gt or _build_opening_line(
        {"name": manual_row["name"], "phone": to_norm},
        role,
    )
    await _prime_opening_audio(camp_id, role, opening_text)
    if not _CAMPAIGN_DATA[camp_id].get("opening_pcm"):
        logger.info(
            "Manual call: Gemini Live will speak the opening on answer"
        )

    from core.state import phone_is_busy, acquire_phone_slot, release_phone_slot, acquire_vobiz_call_slot, release_vobiz_call_slot
    from core.worker import _GLOBAL_CALL_SEMAPHORE, get_next_free_phone_number

    # Automatic failover: if the round-robin pick is busy, switch to the next
    # free outbound line for this role. Only error out when every line is busy.
    if phone_is_busy(from_number):
        logger.warning(
            f"Manual call: phone line {from_number} is busy — trying the next free line."
        )
        free_number = get_next_free_phone_number(role, vobiz_config)
        if free_number and free_number.strip():
            from_number = free_number
            logger.info(f"Manual call: failed over to phone number {from_number} for {role}")
        else:
            raise HTTPException(503, "All phone lines are currently busy. Please try again in a few seconds.")

    global_slot_acquired = False
    slot_acquired = False
    phone_slot_acquired = False

    try:
        await _GLOBAL_CALL_SEMAPHORE.acquire()
        global_slot_acquired = True
        acquire_phone_slot(from_number)
        phone_slot_acquired = True
        
        acquire_vobiz_call_slot(role)
        slot_acquired = True
        
        manual_row["_outbound_phone"] = from_number
        manual_row["_dial_epoch"] = time.time()
        
        auth_tail = auth_id[-6:] if auth_id else ""
        logger.info(
            "Manual Vobiz dial context: role={} auth_id_tail={!r} from_number={!r} base_url={!r} camp_id={}",
            role,
            auth_tail,
            from_number.strip(),
            v_base or "",
            camp_id,
        )
        await make_vobiz_call(
            to=to_norm,
            from_=from_number,
            answer_url=f"{v_base}/vobiz/answer?camp_id={camp_id}",
            hangup_url=f"{v_base}/vobiz/hangup",
            auth_id=auth_id,
            auth_token=auth_token,
        )
        return {"status": "ok", "camp_id": camp_id, "manual_call_id": manual_call_id}
    except Exception as e:
        logger.exception(f"Manual call failed")
        if global_slot_acquired:
            _GLOBAL_CALL_SEMAPHORE.release()
        if phone_slot_acquired:
            release_phone_slot(from_number)
        if slot_acquired:
            release_vobiz_call_slot(role)
        await mark_manual_call_failed(camp_id, str(e))
        _CAMPAIGN_DATA.pop(camp_id, None)
        raise HTTPException(500, str(e))


@router.get("/api/manual/calls/recent")
async def manual_calls_recent(
    request: Request,
    _user: dict = Depends(get_current_user),
    limit: int = 15,
):
    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import list_recent_manual_calls

    rows = await list_recent_manual_calls(role, limit=max(1, min(int(limit), 50)))
    return {"items": [_manual_call_row_to_summary(r) for r in rows]}


@router.get("/api/manual/calls/{call_id}")
async def manual_call_detail(
    call_id: int,
    request: Request,
    _user: dict = Depends(get_current_user),
):
    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import get_manual_call_by_id

    row = await get_manual_call_by_id(call_id)
    if not row or row.get("role") != role:
        raise HTTPException(404, "Manual call not found")
    return _manual_call_detail_response(row)


@router.post("/api/manual/calls/{call_id}/reanalyze")
async def manual_call_reanalyze(
    call_id: int,
    request: Request,
    _user: dict = Depends(get_current_user),
):
    """Re-run post-call Gemini/Gemma QA on the saved JSONL transcript and update SQLite."""
    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import get_manual_call_by_id, update_manual_call_analysis_by_id
    from core.worker import _read_transcript_jsonl
    from services.call_analyzer import analyze_call_transcript

    row = await get_manual_call_by_id(call_id)
    if not row or row.get("role") != role:
        raise HTTPException(404, "Manual call not found")

    log_id = (row.get("log_id") or row.get("camp_id") or "").strip()
    if not log_id:
        raise HTTPException(400, "Call has no log_id transcript yet")

    transcript = ""
    try:
        from services.transcriber import transcribe_audio
        transcribed = await transcribe_audio(log_id, role)
        if transcribed:
            transcript = transcribed
            logger.info("Reanalyze: audio transcription successful for call_id={}", call_id)
    except Exception as e:
        logger.warning("Reanalyze: audio transcription failed for call_id={}: {}", call_id, e)

    if not (transcript or "").strip():
        transcript = _read_transcript_jsonl(role, log_id)
        if (transcript or "").strip():
            logger.info("Reanalyze: falling back to JSONL transcript for call_id={}", call_id)

    if not (transcript or "").strip():
        raise HTTPException(400, "No transcript and no transcribable recording for this call")

    analysis = await analyze_call_transcript(transcript)
    if not await update_manual_call_analysis_by_id(call_id, analysis):
        raise HTTPException(500, "Could not persist analysis")

    refreshed = await get_manual_call_by_id(call_id)
    if not refreshed:
        raise HTTPException(500, "Row missing after update")
    return _manual_call_detail_response(refreshed)


@router.get("/api/manual/calls/{call_id}/recording")
async def manual_call_recording_download(
    call_id: int,
    request: Request,
):
    """Mixed WAV/MP3 with streaming support. Bearer auth or ``?access_token=`` for <audio src>."""
    from loguru import logger
    from core.auth import _decode_jwt
    auth = (request.headers.get("Authorization") or "").strip()
    payload = None
    if auth.startswith("Bearer "):
        payload = _decode_jwt(auth[7:])
        logger.info(f"recording auth: Bearer header, payload={payload}")
    if not payload:
        for key in ("access_token", "token"):
            raw = (request.query_params.get(key) or "").strip()
            logger.info(f"recording auth: trying query key={key}, raw_len={len(raw)}")
            if raw:
                payload = _decode_jwt(raw)
                if payload:
                    logger.info(f"recording auth: query param {key} decoded, payload={payload}")
                    break
                else:
                    logger.warning(f"recording auth: query param {key} failed to decode")
    if not payload:
        logger.warning(f"recording auth: NO payload. auth_header={auth[:30]}, qparams={dict(request.query_params)}")
        raise HTTPException(401, "Not authenticated")

    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import get_manual_call_by_id

    row = await get_manual_call_by_id(call_id)
    if not row or row.get("role") != role:
        raise HTTPException(404, "Manual call not found")
    log_id = (row.get("log_id") or row.get("camp_id") or "").strip()
    if not log_id:
        raise HTTPException(404, "No session log for recording lookup")
    rec = resolve_session_recording_path(log_id)
    if not rec or not rec.is_file():
        raise HTTPException(404, "Recording not found — check CALL_RECORDING_ENABLED and retention")
    media_type = "audio/mpeg" if rec.name.endswith(".mp3") else "audio/wav"
    return range_file_response(rec, request, media_type)


@router.get("/api/incoming/calls/recent")
async def incoming_calls_recent(
    request: Request,
    _user: dict = Depends(get_current_user),
    limit: int = 15,
):
    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import list_recent_incoming_calls

    rows = await list_recent_incoming_calls(role, limit=max(1, min(int(limit), 5000)))
    return {"items": [_incoming_call_row_to_summary(r) for r in rows]}


@router.get("/api/incoming/calls/{call_id}")
async def incoming_call_detail(
    call_id: int,
    request: Request,
    _user: dict = Depends(get_current_user),
):
    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import get_incoming_call_by_id

    row = await get_incoming_call_by_id(call_id)
    if not row or row.get("role") != role:
        raise HTTPException(404, "Incoming call not found")
    return _incoming_call_detail_response(row)


@router.post("/api/incoming/calls/{call_id}/reanalyze")
async def incoming_call_reanalyze(
    call_id: int,
    request: Request,
    _user: dict = Depends(get_current_user),
):
    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import get_incoming_call_by_id, update_incoming_call_analysis_by_id
    from core.worker import _read_transcript_jsonl
    from services.call_analyzer import analyze_call_transcript

    row = await get_incoming_call_by_id(call_id)
    if not row or row.get("role") != role:
        raise HTTPException(404, "Incoming call not found")

    log_id = (row.get("log_id") or row.get("camp_id") or "").strip()
    if not log_id:
        raise HTTPException(400, "Call has no log_id transcript yet")

    transcript = ""
    try:
        from services.transcriber import transcribe_audio
        transcribed = await transcribe_audio(log_id, role)
        if transcribed:
            transcript = transcribed
            logger.info("Incoming reanalyze: audio transcription successful for call_id={}", call_id)
    except Exception as e:
        logger.warning("Incoming reanalyze: audio transcription failed for call_id={}: {}", call_id, e)

    if not (transcript or "").strip():
        transcript = _read_transcript_jsonl(role, log_id)
        if (transcript or "").strip():
            logger.info("Incoming reanalyze: falling back to JSONL transcript for call_id={}", call_id)

    if not (transcript or "").strip():
        raise HTTPException(400, "No transcript and no transcribable recording for this call")

    analysis = await analyze_call_transcript(transcript)
    if not await update_incoming_call_analysis_by_id(call_id, analysis):
        raise HTTPException(500, "Could not persist analysis")

    refreshed = await get_incoming_call_by_id(call_id)
    if not refreshed:
        raise HTTPException(500, "Row missing after update")
    return _incoming_call_detail_response(refreshed)


@router.get("/api/incoming/calls/{call_id}/recording")
async def incoming_call_recording_download(
    call_id: int,
    request: Request,
):
    from loguru import logger
    from core.auth import _decode_jwt
    auth = (request.headers.get("Authorization") or "").strip()
    payload = None
    if auth.startswith("Bearer "):
        payload = _decode_jwt(auth[7:])
    if not payload:
        for key in ("access_token", "token"):
            raw = (request.query_params.get(key) or "").strip()
            if raw:
                payload = _decode_jwt(raw)
                if payload:
                    break
    if not payload:
        raise HTTPException(401, "Not authenticated")

    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )
    from core.storage import get_incoming_call_by_id

    row = await get_incoming_call_by_id(call_id)
    if not row or row.get("role") != role:
        raise HTTPException(404, "Incoming call not found")
    log_id = (row.get("log_id") or row.get("camp_id") or "").strip()
    if not log_id:
        raise HTTPException(404, "No session log for recording lookup")
    rec = resolve_session_recording_path(log_id)
    if not rec or not rec.is_file():
        raise HTTPException(404, "Recording not found — check CALL_RECORDING_ENABLED and retention")
    media_type = "audio/mpeg" if rec.name.endswith(".mp3") else "audio/wav"
    return range_file_response(rec, request, media_type)


class RescheduleCampaignReq(BaseModel):
    from_date: str
    to_date: str
    outcomes: list[str]
    target_datetime: str


@router.post("/api/campaign/reschedule")
async def reschedule_campaign_calls(
    data: RescheduleCampaignReq,
    request: Request,
    _user: dict = Depends(get_current_user),
):
    """Reschedule historical campaign leads to a future callback datetime."""
    from datetime import datetime
    from core.storage import reschedule_leads

    role = normalize_console_role(
        request.query_params.get("role") or request.headers.get("X-User-Role") or "sales_1"
    )

    # Parse target datetime (ISO 8601) to epoch seconds.
    target_dt_str = (data.target_datetime or "").strip()
    if not target_dt_str:
        raise HTTPException(400, "Target date/time is required")
    try:
        # Handle both "2026-06-13T18:00" and "2026-06-13T18:00:00" plus offsets.
        if target_dt_str.endswith("Z"):
            target_dt_str = target_dt_str[:-1] + "+00:00"
        target_dt = datetime.fromisoformat(target_dt_str)
        if target_dt.tzinfo is None:
            # Assume Asia/Kolkata when no timezone is provided (matches dashboard UI).
            from zoneinfo import ZoneInfo
            target_dt = target_dt.replace(tzinfo=ZoneInfo("Asia/Kolkata"))
        target_epoch = target_dt.timestamp()
    except Exception as exc:
        logger.warning("Invalid target_datetime {}: {}", data.target_datetime, exc)
        raise HTTPException(400, f"Invalid target date/time: {data.target_datetime}")

    if target_epoch <= time.time():
        raise HTTPException(400, "Target date/time must be in the future")

    valid_outcomes = {"failed_no_answer", "interested", "cut_in_middle", "not_interested"}
    outcomes = [o for o in data.outcomes if o in valid_outcomes]
    if not outcomes:
        raise HTTPException(400, "Select at least one outcome to reschedule")

    try:
        count = await reschedule_leads(
            role,
            data.from_date,
            data.to_date,
            outcomes,
            target_epoch,
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        logger.exception("Reschedule campaign failed")
        raise HTTPException(500, f"Reschedule failed: {exc}")

    return {"status": "ok", "rescheduled_count": count}


@router.get("/api/conversation-logs/{date}/{log_id}")
async def get_conversation_log(date: str, log_id: str):
    log_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "logs", date, f"{log_id}.txt")
    if os.path.exists(log_path):
        return FileResponse(log_path)
    raise HTTPException(404, "Log not found")

@router.get("/api/recordings/{date}/{filename}")
async def get_recording(date: str, filename: str, request: Request):
    rec_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "logs", date, filename)
    if os.path.exists(rec_path):
        media_type = "audio/mpeg" if filename.endswith(".mp3") else "audio/wav"
        return range_file_response(Path(rec_path), request, media_type)
    raise HTTPException(404, "Recording not found")



# ── Virtual Meet API ───────────────────────────────────────────────


@router.get("/api/campaign/lead/{lead_id}/virtual-meet")
async def get_virtual_meet(lead_id: int, request: Request):
    """Return the latest virtual meet record for a lead."""
    from core.storage import get_virtual_meet_for_lead

    role = extract_role(request)
    vm = await get_virtual_meet_for_lead(lead_id)
    if vm and vm["role"] == role:
        return vm
    return {"id": None}


@router.post("/api/campaign/lead/{lead_id}/virtual-meet/reschedule")
async def reschedule_virtual_meet(lead_id: int, request: Request):
    """Reschedule a virtual meet for a lead."""
    from core.storage import (
        get_virtual_meet_for_lead,
        reschedule_virtual_meet as _reschedule_vm,
    )

    role = extract_role(request)
    body = await request.json()
    new_date = (body.get("meet_date") or "").strip()
    new_time = (body.get("meet_time") or "").strip()
    new_notes = (body.get("notes") or "").strip()
    if not new_date or not new_time:
        raise HTTPException(400, "meet_date and meet_time are required")

    vm = await get_virtual_meet_for_lead(lead_id)
    if not vm:
        # No existing meet — create a new one
        from core.storage import add_virtual_meet as _add_vm

        new_id = await _add_vm(lead_id, role, new_date, new_time, new_notes)
        return {"status": "ok", "id": new_id, "action": "created"}

    ok = await _reschedule_vm(vm["id"], new_date, new_time, new_notes)
    if not ok:
        raise HTTPException(500, "Failed to reschedule virtual meet")
    return {"status": "ok", "id": vm["id"], "action": "rescheduled"}


@router.post("/api/campaign/lead/{lead_id}/virtual-meet/cancel")
async def cancel_virtual_meet(lead_id: int, request: Request):
    """Cancel a virtual meet for a lead."""
    from core.storage import get_virtual_meet_for_lead, cancel_virtual_meet as _cancel_vm

    role = extract_role(request)
    vm = await get_virtual_meet_for_lead(lead_id)
    if not vm or vm["role"] != role:
        raise HTTPException(404, "No virtual meet found for this lead")
    ok = await _cancel_vm(vm["id"])
    if not ok:
        raise HTTPException(500, "Failed to cancel virtual meet")
    return {"status": "ok"}
